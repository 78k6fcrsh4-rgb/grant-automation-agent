from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
from openpyxl.utils import get_column_letter
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from icalendar import Calendar, Event, Alarm
from datetime import datetime, timedelta
from app.models.schemas import GrantData
from typing import Dict, List
import os
import re


class DocumentService:
    """Universal document generation service that adapts to any grant format"""
    
    def __init__(self, temp_dir: str = "temp_files"):
        self.temp_dir = temp_dir
        os.makedirs(temp_dir, exist_ok=True)
    
    def generate_workplan_pdf(self, grant_data: GrantData, file_id: str) -> str:
        """Generate comprehensive work plan PDF adapted to grant structure"""
        filename = f"{file_id}_workplan.pdf"
        filepath = os.path.join(self.temp_dir, filename)
        
        doc = SimpleDocTemplate(
            filepath, 
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=1*inch,
            bottomMargin=0.75*inch
        )
        
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.HexColor('#1a365d'),
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        subtitle_style = ParagraphStyle(
            'Subtitle',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.HexColor('#4a5568'),
            spaceAfter=20,
            alignment=TA_CENTER
        )
        
        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#2d3748'),
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        # Title Page
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph("GRANT IMPLEMENTATION WORK PLAN", title_style))
        story.append(Spacer(1, 0.1*inch))
        
        grant_title = grant_data.grant_title or "Grant Project"
        story.append(Paragraph(f"<b>{grant_title}</b>", subtitle_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Grant Overview Box
        overview_data = [
            ['GRANT OVERVIEW'],
            ['Recipient Organization:', grant_data.organization_name or 'N/A'],
            ['Funding Agency:', grant_data.funder_name or 'N/A'],
            ['Grant Period:', grant_data.grant_period or 'N/A'],
            ['Total Award Amount:', f"${grant_data.grant_amount:,.2f}" if grant_data.grant_amount else 'N/A'],
            ['Document Prepared:', datetime.now().strftime('%B %d, %Y')],
        ]
        
        overview_table = Table(overview_data, colWidths=[2*inch, 4*inch])
        overview_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c5282')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('SPAN', (0, 0), (-1, 0)),
            ('BACKGROUND', (0, 1), (0, -1), colors.HexColor('#edf2f7')),
            ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e0')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        
        story.append(overview_table)
        story.append(PageBreak())
        
        # Work Plan Tasks
        if grant_data.workplan and grant_data.workplan.tasks:
            story.append(Paragraph("IMPLEMENTATION PLAN", heading2_style))
            story.append(Spacer(1, 0.2*inch))
            
            for idx, task in enumerate(grant_data.workplan.tasks, 1):
                # Task header
                task_header_data = [[f"ACTIVITY {idx}: {task.task_name.upper()}"]]
                task_header_table = Table(task_header_data, colWidths=[6.5*inch])
                task_header_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#4299e1')),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.white),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 11),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 12),
                    ('TOPPADDING', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                ]))
                story.append(task_header_table)
                
                # Task details
                task_details_data = [
                    ['Description:', task.description or 'No description provided'],
                    ['Timeline:', f"{task.start_date or 'TBD'} to {task.end_date or 'TBD'}"],
                    ['Responsible Party:', task.responsible_party or 'To be assigned'],
                    ['Deliverables:', task.deliverables or 'See description'],
                ]
                
                task_details_table = Table(task_details_data, colWidths=[1.5*inch, 5*inch])
                task_details_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e2e8f0')),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                    ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e0')),
                ]))
                
                story.append(task_details_table)
                story.append(Spacer(1, 0.3*inch))
        
        # Timeline & Milestones
        if grant_data.timeline and grant_data.timeline.items:
            story.append(PageBreak())
            story.append(Paragraph("PROJECT TIMELINE & MILESTONES", heading2_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Group timeline items by category
            timeline_by_category = {}
            for item in grant_data.timeline.items:
                category = item.category or 'general'
                if category not in timeline_by_category:
                    timeline_by_category[category] = []
                timeline_by_category[category].append(item)
            
            # Display by category
            for category, items in timeline_by_category.items():
                story.append(Paragraph(f"{category.upper()} DEADLINES", ParagraphStyle(
                    'CategoryHeader',
                    parent=styles['Heading3'],
                    fontSize=12,
                    textColor=colors.HexColor('#2c5282'),
                    spaceAfter=8,
                    spaceBefore=8,
                    fontName='Helvetica-Bold'
                )))
                
                timeline_data = [['Date', 'Description', 'Amount']]
                
                for item in sorted(items, key=lambda x: self._parse_date_safe(x.date)):
                    timeline_data.append([
                        item.date,
                        item.description,
                        item.amount or '-'
                    ])
                
                timeline_table = Table(timeline_data, colWidths=[1.2*inch, 4*inch, 1.3*inch])
                timeline_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c5282')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('ALIGN', (0, 1), (0, -1), 'CENTER'),
                    ('ALIGN', (1, 1), (1, -1), 'LEFT'),
                    ('ALIGN', (2, 1), (2, -1), 'RIGHT'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e0')),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ]))
                
                story.append(timeline_table)
                story.append(Spacer(1, 0.2*inch))
        
        doc.build(story)
        return filepath
    
    def generate_budget_excel(self, grant_data: GrantData, file_id: str) -> str:
        """
        Generate comprehensive budget Excel matching government grant format
        Creates: 1) Budget Summary 2) Reimbursement Tracker 3) Disbursement Schedule
        """
        filename = f"{file_id}_budget.xlsx"
        filepath = os.path.join(self.temp_dir, filename)
        
        wb = Workbook()
        
        # Remove default sheet
        if 'Sheet' in wb.sheetnames:
            wb.remove(wb['Sheet'])
        
        # Style definitions
        header_fill = PatternFill(start_color="1a365d", end_color="1a365d", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True, size=11)
        yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        yellow_font = Font(bold=True, size=10)
        subheader_fill = PatternFill(start_color="2c5282", end_color="2c5282", fill_type="solid")
        subheader_font = Font(color="FFFFFF", bold=True, size=10)
        total_fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")
        total_font = Font(bold=True, size=11)
        currency_format = '"$"#,##0.00'
        border_thin = Border(
            left=Side(style='thin', color='000000'),
            right=Side(style='thin', color='000000'),
            top=Side(style='thin', color='000000'),
            bottom=Side(style='thin', color='000000')
        )
        
        # ============================================
        # SHEET 1: Budget Summary (Government Format)
        # ============================================
        ws_budget = wb.create_sheet("Budget Summary")
        
        # Title section with yellow highlighting (like Chicago format)
        ws_budget['A1'] = 'GRANT BUDGET SUMMARY'
        ws_budget['A1'].font = Font(bold=True, size=14)
        ws_budget['A1'].fill = yellow_fill
        ws_budget.merge_cells('A1:E1')
        
        # Grant identification
        row = 3
        info_fields = [
            ('Delegate Agency:', grant_data.organization_name or '[Organization Name]'),
            ('Program Name:', grant_data.grant_title or '[Grant Program Name]'),
            ('Funding Agency:', grant_data.funder_name or '[Funding Agency]'),
            ('Grant Period:', grant_data.grant_period or '[Start Date] - [End Date]'),
            ('Total Award:', grant_data.grant_amount or 0),
        ]
        
        for label, value in info_fields:
            ws_budget[f'A{row}'] = label
            ws_budget[f'A{row}'].font = Font(bold=True)
            ws_budget[f'A{row}'].fill = yellow_fill
            ws_budget[f'B{row}'] = value
            if 'Award' in label and isinstance(value, (int, float)):
                ws_budget[f'B{row}'].number_format = currency_format
            row += 1
        
        row += 1
        ws_budget[f'A{row}'] = 'Note: The entire budget for this program must be shown.'
        ws_budget[f'A{row}'].font = Font(italic=True, size=9)
        ws_budget.merge_cells(f'A{row}:E{row}')
        
        row += 2
        
        # Budget table headers
        headers = ['Item of Expenditure', 'Account #', 'Grant Share', 'Other Share', 'Total Cost']
        for col_num, header in enumerate(headers, 1):
            cell = ws_budget.cell(row=row, column=col_num)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = border_thin
        
        header_row = row
        row += 1
        
        # Budget line items
        total_grant_share = 0
        total_other_share = 0
        
        if grant_data.budget and grant_data.budget.items:
            for item in grant_data.budget.items:
                ws_budget.cell(row=row, column=1, value=item.category).border = border_thin
                ws_budget.cell(row=row, column=2, value=self._extract_account_number(item.category)).border = border_thin
                
                # For this example, allocate 95% to grant, 5% to other (adjust as needed)
                grant_share = item.amount * 0.95
                other_share = item.amount * 0.05
                
                grant_cell = ws_budget.cell(row=row, column=3, value=grant_share)
                grant_cell.number_format = currency_format
                grant_cell.border = border_thin
                
                other_cell = ws_budget.cell(row=row, column=4, value=other_share)
                other_cell.number_format = currency_format
                other_cell.border = border_thin
                
                total_cell = ws_budget.cell(row=row, column=5, value=item.amount)
                total_cell.number_format = currency_format
                total_cell.border = border_thin
                
                total_grant_share += grant_share
                total_other_share += other_share
                row += 1
        
        # Totals row
        ws_budget.cell(row=row, column=1, value='TOTALS').font = total_font
        ws_budget.cell(row=row, column=1).fill = total_fill
        ws_budget.cell(row=row, column=1).border = border_thin
        
        ws_budget.cell(row=row, column=2, value='').fill = total_fill
        ws_budget.cell(row=row, column=2).border = border_thin
        
        for col in [3, 4, 5]:
            cell = ws_budget.cell(row=row, column=col)
            if col == 3:
                cell.value = total_grant_share
            elif col == 4:
                cell.value = total_other_share
            else:
                cell.value = total_grant_share + total_other_share
            cell.number_format = currency_format
            cell.font = total_font
            cell.fill = total_fill
            cell.border = border_thin
        
        row += 1
        ws_budget.cell(row=row, column=1, value='***ALL COLUMNS / ROWS MUST BALANCE***').font = Font(bold=True, size=9)
        ws_budget.merge_cells(f'A{row}:E{row}')
        
        # Column widths
        ws_budget.column_dimensions['A'].width = 30
        ws_budget.column_dimensions['B'].width = 12
        ws_budget.column_dimensions['C'].width = 15
        ws_budget.column_dimensions['D'].width = 15
        ws_budget.column_dimensions['E'].width = 15
        
        # ============================================
        # SHEET 2: Reimbursement Request Tracker
        # ============================================
        ws_reimburse = wb.create_sheet("Reimbursement Tracker")
        
        ws_reimburse['A1'] = 'REIMBURSEMENT REQUEST TRACKER'
        ws_reimburse['A1'].font = Font(bold=True, size=14)
        ws_reimburse.merge_cells('A1:H1')
        
        ws_reimburse['A3'] = 'Track all reimbursement requests submitted to the funding agency'
        ws_reimburse.merge_cells('A3:H3')
        
        # Headers
        reimburse_headers = [
            'Request #',
            'Submission Date',
            'Period Covered',
            'Amount Requested',
            'Amount Approved',
            'Payment Date',
            'Status',
            'Notes'
        ]
        
        for col_num, header in enumerate(reimburse_headers, 1):
            cell = ws_reimburse.cell(row=5, column=col_num)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = border_thin
        
        # Sample rows for tracking
        for row_num in range(6, 16):
            ws_reimburse.cell(row=row_num, column=1, value=row_num - 5).border = border_thin
            for col in range(2, 9):
                cell = ws_reimburse.cell(row=row_num, column=col, value='')
                cell.border = border_thin
                if col == 4 or col == 5:
                    cell.number_format = currency_format
                cell.fill = PatternFill(start_color="fef5e7", end_color="fef5e7", fill_type="solid")
        
        # Column widths
        ws_reimburse.column_dimensions['A'].width = 10
        ws_reimburse.column_dimensions['B'].width = 15
        ws_reimburse.column_dimensions['C'].width = 20
        ws_reimburse.column_dimensions['D'].width = 15
        ws_reimburse.column_dimensions['E'].width = 15
        ws_reimburse.column_dimensions['F'].width = 15
        ws_reimburse.column_dimensions['G'].width = 15
        ws_reimburse.column_dimensions['H'].width = 30
        
        # ============================================
        # SHEET 3: Disbursement Schedule
        # ============================================
        ws_disbursement = wb.create_sheet("Disbursement Schedule")
        
        ws_disbursement['A1'] = 'GRANT DISBURSEMENT SCHEDULE'
        ws_disbursement['A1'].font = Font(bold=True, size=14)
        ws_disbursement.merge_cells('A1:G1')
        
        ws_disbursement['A3'] = 'Expected and actual payment receipts from funding agency'
        ws_disbursement.merge_cells('A3:G3')
        
        # Headers
        disburse_headers = [
            'Payment #',
            'Expected Date',
            'Expected Amount',
            'Actual Date',
            'Actual Amount',
            'Status',
            'Notes'
        ]
        
        for col_num, header in enumerate(disburse_headers, 1):
            cell = ws_disbursement.cell(row=5, column=col_num)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = border_thin
        
        # Add payment schedule from timeline
        row = 6
        payment_num = 1
        if grant_data.timeline and grant_data.timeline.items:
            for item in sorted(grant_data.timeline.items, key=lambda x: self._parse_date_safe(x.date)):
                if item.category == 'payment' and item.amount:
                    ws_disbursement.cell(row=row, column=1, value=payment_num).border = border_thin
                    ws_disbursement.cell(row=row, column=2, value=item.date).border = border_thin
                    
                    # Extract amount
                    amount = self._extract_amount(item.amount)
                    amount_cell = ws_disbursement.cell(row=row, column=3, value=amount)
                    amount_cell.number_format = currency_format
                    amount_cell.border = border_thin
                    
                    # Actual columns (for user to fill)
                    ws_disbursement.cell(row=row, column=4, value='').border = border_thin
                    ws_disbursement.cell(row=row, column=4).fill = PatternFill(start_color="fef5e7", end_color="fef5e7", fill_type="solid")
                    
                    actual_amount_cell = ws_disbursement.cell(row=row, column=5, value=0)
                    actual_amount_cell.number_format = currency_format
                    actual_amount_cell.border = border_thin
                    actual_amount_cell.fill = PatternFill(start_color="fef5e7", end_color="fef5e7", fill_type="solid")
                    
                    status_cell = ws_disbursement.cell(row=row, column=6, value='Pending')
                    status_cell.border = border_thin
                    status_cell.fill = PatternFill(start_color="fef5e7", end_color="fef5e7", fill_type="solid")
                    
                    notes_cell = ws_disbursement.cell(row=row, column=7, value=item.description)
                    notes_cell.border = border_thin
                    
                    row += 1
                    payment_num += 1
        
        # Column widths
        for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
            ws_disbursement.column_dimensions[col].width = 18
        
        # ============================================
        # SHEET 4: Expense Tracking Log
        # ============================================
        ws_expenses = wb.create_sheet("Expense Log")
        
        ws_expenses['A1'] = 'DETAILED EXPENSE TRACKING LOG'
        ws_expenses['A1'].font = Font(bold=True, size=14)
        ws_expenses.merge_cells('A1:I1')
        
        ws_expenses['A3'] = 'Record all expenses charged to this grant'
        ws_expenses.merge_cells('A3:I3')
        
        # Headers
        expense_headers = [
            'Date',
            'Vendor/Payee',
            'Description',
            'Budget Category',
            'Account #',
            'Amount',
            'Payment Method',
            'Receipt/Invoice #',
            'Notes'
        ]
        
        for col_num, header in enumerate(expense_headers, 1):
            cell = ws_expenses.cell(row=5, column=col_num)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = border_thin
        
        # Add sample row
        ws_expenses.cell(row=6, column=1, value=datetime.now().strftime('%Y-%m-%d')).border = border_thin
        ws_expenses.cell(row=6, column=2, value='Example Vendor').border = border_thin
        ws_expenses.cell(row=6, column=3, value='Sample expense description').border = border_thin
        ws_expenses.cell(row=6, column=4, value='Personnel').border = border_thin
        ws_expenses.cell(row=6, column=5, value='0005').border = border_thin
        
        sample_amount = ws_expenses.cell(row=6, column=6, value=0)
        sample_amount.number_format = currency_format
        sample_amount.border = border_thin
        sample_amount.fill = PatternFill(start_color="fef5e7", end_color="fef5e7", fill_type="solid")
        
        for col in range(7, 10):
            cell = ws_expenses.cell(row=6, column=col, value='')
            cell.border = border_thin
            cell.fill = PatternFill(start_color="fef5e7", end_color="fef5e7", fill_type="solid")
        
        # Column widths
        ws_expenses.column_dimensions['A'].width = 12
        ws_expenses.column_dimensions['B'].width = 25
        ws_expenses.column_dimensions['C'].width = 35
        ws_expenses.column_dimensions['D'].width = 20
        ws_expenses.column_dimensions['E'].width = 12
        ws_expenses.column_dimensions['F'].width = 15
        ws_expenses.column_dimensions['G'].width = 15
        ws_expenses.column_dimensions['H'].width = 15
        ws_expenses.column_dimensions['I'].width = 30
        
        wb.save(filepath)
        return filepath
    
    def generate_report_template_docx(self, grant_data: GrantData, file_id: str) -> str:
        """Generate progress report template matching grant requirements"""
        filename = f"{file_id}_report_template.docx"
        filepath = os.path.join(self.temp_dir, filename)
        
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Helper function
        def add_colored_heading(text, level=1, color='1a365d'):
            heading = doc.add_heading(text, level)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in heading.runs:
                run.font.color.rgb = RGBColor(
                    int(color[0:2], 16),
                    int(color[2:4], 16),
                    int(color[4:6], 16)
                )
            return heading
        
        # Title Page
        title = doc.add_heading('Grant Progress Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(26, 54, 93)
        
        doc.add_paragraph()
        
        subtitle = doc.add_paragraph(grant_data.grant_title or '[Grant Program Name]')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.bold = True
        subtitle_run.font.color.rgb = RGBColor(74, 85, 104)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Grant Information
        add_colored_heading('Grant Information', 1)
        
        info_table = doc.add_table(rows=7, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('Recipient Organization:', grant_data.organization_name or '[Organization Name]'),
            ('Grant Program:', grant_data.grant_title or '[Grant Program Name]'),
            ('Funding Agency:', grant_data.funder_name or '[Funding Agency]'),
            ('Grant Period:', grant_data.grant_period or '[Grant Period]'),
            ('Total Award Amount:', f"${grant_data.grant_amount:,.2f}" if grant_data.grant_amount else '[Amount]'),
            ('Reporting Period:', '[Start Date] to [End Date]'),
            ('Report Submission Date:', datetime.now().strftime('%B %d, %Y')),
        ]
        
        for idx, (label, value) in enumerate(info_data):
            info_table.rows[idx].cells[0].text = label
            info_table.rows[idx].cells[1].text = value
            info_table.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True
            
            # Shading
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'edf2f7')
            info_table.rows[idx].cells[0]._element.get_or_add_tcPr().append(shading_elm)
        
        doc.add_paragraph()
        doc.add_page_break()
        
        # Executive Summary
        add_colored_heading('Executive Summary', 1)
        doc.add_paragraph('[Provide a brief overview of activities, achievements, and overall progress during this reporting period.]')
        doc.add_paragraph()
        
        # Progress by Activity
        add_colored_heading('Progress by Activity', 1)
        
        if grant_data.workplan and grant_data.workplan.tasks:
            for idx, task in enumerate(grant_data.workplan.tasks, 1):
                add_colored_heading(f'Activity {idx}: {task.task_name}', 2, '2d3748')
                
                doc.add_paragraph(f'Planned Activity: {task.description}')
                doc.add_paragraph(f'Expected Deliverables: {task.deliverables or "See description"}')
                doc.add_paragraph()
                
                doc.add_paragraph('Progress Made:', style='Heading 3')
                doc.add_paragraph('[Describe specific progress on this activity]', style='List Bullet')
                doc.add_paragraph('[Include measurable results and outcomes]', style='List Bullet')
                doc.add_paragraph()
                
                doc.add_paragraph('Challenges:', style='Heading 3')
                doc.add_paragraph('[Note any obstacles encountered and how they were addressed]', style='List Bullet')
                doc.add_paragraph()
        
        # Financial Report
        doc.add_page_break()
        add_colored_heading('Financial Report', 1)
        
        if grant_data.budget and grant_data.budget.items:
            financial_table = doc.add_table(rows=len(grant_data.budget.items) + 2, cols=5)
            financial_table.style = 'Light Grid Accent 1'
            
            # Headers
            headers = ['Budget Category', 'Budgeted', 'Spent to Date', 'Remaining', '% Utilized']
            header_cells = financial_table.rows[0].cells
            for idx, header in enumerate(headers):
                header_cells[idx].text = header
                header_cells[idx].paragraphs[0].runs[0].font.bold = True
                
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '2c5282')
                header_cells[idx]._element.get_or_add_tcPr().append(shading_elm)
                header_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            
            # Data rows
            for idx, item in enumerate(grant_data.budget.items, 1):
                row_cells = financial_table.rows[idx].cells
                row_cells[0].text = item.category
                row_cells[1].text = f"${item.amount:,.2f}"
                row_cells[2].text = '$[Amount Spent]'
                row_cells[3].text = '$[Remaining]'
                row_cells[4].text = '[%]'
            
            # Total row
            total_cells = financial_table.rows[-1].cells
            total_cells[0].text = 'TOTAL'
            total_cells[0].paragraphs[0].runs[0].font.bold = True
            total_cells[1].text = f"${grant_data.budget.total_grant_amount:,.2f}"
            total_cells[1].paragraphs[0].runs[0].font.bold = True
            total_cells[2].text = '$[Total Spent]'
            total_cells[2].paragraphs[0].runs[0].font.bold = True
            total_cells[3].text = '$[Total Remaining]'
            total_cells[3].paragraphs[0].runs[0].font.bold = True
            total_cells[4].text = '[%]'
            total_cells[4].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Outcomes and Impact
        doc.add_page_break()
        add_colored_heading('Outcomes and Impact', 1)
        doc.add_paragraph('[Describe measurable outcomes and impact achieved during this period]')
        doc.add_paragraph()
        
        # Attachments
        add_colored_heading('Supporting Documentation', 1)
        doc.add_paragraph('☐ Financial statements', style='List Bullet')
        doc.add_paragraph('☐ Receipts and invoices', style='List Bullet')
        doc.add_paragraph('☐ Photos or media', style='List Bullet')
        doc.add_paragraph('☐ Participant feedback', style='List Bullet')
        doc.add_paragraph('☐ Other: [Specify]', style='List Bullet')
        doc.add_paragraph()
        
        # Signature
        doc.add_paragraph()
        sig_table = doc.add_table(rows=3, cols=2)
        sig_table.rows[0].cells[0].text = 'Prepared by:'
        sig_table.rows[0].cells[1].text = 'Date:'
        sig_table.rows[1].cells[0].text = '\n_______________________________'
        sig_table.rows[1].cells[1].text = '\n_______________________________'
        sig_table.rows[2].cells[0].text = '[Name and Title]'
        
        doc.save(filepath)
        return filepath
    
    def generate_agenda_template_docx(self, grant_data: GrantData, file_id: str) -> str:
        """Generate a status meeting agenda template."""
        filename = f"{file_id}_agenda.docx"
        filepath = os.path.join(self.temp_dir, filename)

        doc = Document()
        doc.add_heading('Grant Status Meeting Agenda', 0)
        doc.add_paragraph(f"Grant: {grant_data.grant_title or '[Grant Title]'}")
        doc.add_paragraph(f"Organization: {grant_data.organization_name or '[Organization]'}")
        doc.add_paragraph(f"Funder: {grant_data.funder_name or '[Funder]'}")
        doc.add_paragraph(f"Grant Period: {grant_data.grant_period or '[Grant Period]'}")
        doc.add_paragraph()
        doc.add_heading('Standing Agenda', level=1)
        for item in [
            '1. Review action items from prior meeting',
            '2. Progress against work plan milestones',
            '3. Upcoming deliverables and deadlines',
            '4. Reimbursement / disbursement status and variance',
            '5. Reporting requirements and submission readiness',
            '6. Risks, dependencies, and decisions needed',
            '7. Next steps and owners',
        ]:
            doc.add_paragraph(item, style='List Bullet')
        doc.add_paragraph()
        doc.add_heading('Upcoming Key Dates', level=1)
        if grant_data.timeline and grant_data.timeline.items:
            for tl_item in grant_data.timeline.items[:12]:
                doc.add_paragraph(f"{tl_item.date}: {tl_item.description}", style='List Bullet')
        else:
            doc.add_paragraph('[Populate from grant timeline]')
        doc.save(filepath)
        return filepath

    def _build_event(self, *, summary: str, description: str, start_dt: datetime, end_dt: datetime, all_day: bool = False, alarms: List[timedelta] | None = None, uid: str) -> Event:
        event = Event()
        event.add('summary', summary)
        if all_day:
            event.add('dtstart', start_dt.date())
            event.add('dtend', end_dt.date())
        else:
            event.add('dtstart', start_dt)
            event.add('dtend', end_dt)
        event.add('description', description)
        event.add('dtstamp', datetime.now())
        event.add('uid', uid)
        event.add('status', 'CONFIRMED')
        for trigger in alarms or []:
            alarm = Alarm()
            alarm.add('action', 'DISPLAY')
            alarm.add('description', summary)
            alarm.add('trigger', trigger)
            event.add_component(alarm)
        return event

    def _grant_period_bounds(self, grant_data: GrantData):
        dates = []
        if grant_data.timeline and grant_data.timeline.items:
            for item in grant_data.timeline.items:
                dt = self._parse_date_safe(item.date)
                if dt.year < 2099:
                    dates.append(dt)
        if dates:
            return min(dates), max(dates)
        now = datetime.now()
        return now, now + timedelta(days=180)

    def _detect_calendar_discrepancies(
        self,
        grant_data: GrantData,
        disbursement_reminder_days: int,
    ) -> List[str]:
        """Detect mismatches between disbursement timeline events and submission requirements."""
        discrepancies: List[str] = []

        disbursement_events = []
        if grant_data.timeline and grant_data.timeline.items:
            for item in grant_data.timeline.items:
                cat = (item.category or '').lower()
                if cat in ('disbursement', 'reimbursement'):
                    dt = self._parse_date_safe(item.date)
                    if dt.year < 2099:
                        disbursement_events.append((dt, item.description))

        req_dates = []
        for req in grant_data.submission_requirements or []:
            if req.due_date:
                dt = self._parse_date_safe(req.due_date)
                if dt.year < 2099:
                    req_dates.append((dt, req.category))

        # Each submission requirement should have a matching disbursement event nearby
        for req_dt, req_cat in req_dates:
            matched = any(
                abs((d - req_dt).days) <= disbursement_reminder_days
                for d, _ in disbursement_events
            )
            if not matched and disbursement_events:
                discrepancies.append(
                    f"Submission requirement '{req_cat}' due {req_dt.strftime('%Y-%m-%d')} "
                    f"has no disbursement event within {disbursement_reminder_days} days."
                )

        # Each disbursement event should have a matching submission requirement nearby
        for d_dt, d_desc in disbursement_events:
            matched = any(
                abs((req_dt - d_dt).days) <= disbursement_reminder_days
                for req_dt, _ in req_dates
            )
            if not matched and req_dates:
                discrepancies.append(
                    f"Disbursement event '{d_desc}' on {d_dt.strftime('%Y-%m-%d')} "
                    f"has no matching submission requirement within {disbursement_reminder_days} days."
                )

        return discrepancies

    def generate_meeting_calendar_ics(
        self,
        grant_data: GrantData,
        file_id: str,
        meeting_interval_days: int = 14,
    ) -> str:
        """Generate a recurring status-meeting calendar (bi-weekly by default).

        Each event includes a standard agenda template in the DESCRIPTION.
        Uses RRULE for clean recurrence rather than individual events.
        """
        filename = f"{file_id}_meeting_calendar.ics"
        filepath = os.path.join(self.temp_dir, filename)
        cal = Calendar()
        cal.add('prodid', '-//Grant Status Meetings//EN//')
        cal.add('version', '2.0')
        cal.add('calscale', 'GREGORIAN')
        cal.add('method', 'PUBLISH')
        cal.add('x-wr-calname', f"Meetings: {grant_data.grant_title or 'Grant'}")
        cal.add('x-wr-timezone', 'America/Chicago')

        start_bound, end_bound = self._grant_period_bounds(grant_data)
        first_meeting = start_bound.replace(hour=10, minute=0, second=0, microsecond=0)

        title = grant_data.grant_title or "Grant"
        agenda_description = (
            f"Status meeting: {title}\n"
            f"Cadence: every {meeting_interval_days} days\n\n"
            "AGENDA TEMPLATE\n"
            "1. Progress against milestones and deliverables\n"
            "2. Budget vs. actuals review\n"
            "3. Upcoming reporting deadlines\n"
            "4. Upcoming disbursement requests\n"
            "5. Issues, risks, and action items\n"
            "6. Next steps and owner assignments"
        )

        from icalendar import vRecur
        event = Event()
        event.add('summary', f"Status Meeting: {title}")
        event.add('dtstart', first_meeting)
        event.add('dtend', first_meeting + timedelta(minutes=50))
        event.add('description', agenda_description)
        event.add('dtstamp', datetime.now())
        event.add('uid', f"{file_id}-meeting-recurring@grantmanagement.local")
        event.add('status', 'CONFIRMED')
        # RRULE: repeat every N days until the end of the grant period
        event.add('rrule', vRecur(freq='DAILY', interval=meeting_interval_days, until=end_bound))
        # 10-minute reminder
        alarm = Alarm()
        alarm.add('action', 'DISPLAY')
        alarm.add('description', f"Starting soon: Status Meeting — {title}")
        alarm.add('trigger', timedelta(minutes=-10))
        event.add_component(alarm)
        cal.add_component(event)

        with open(filepath, 'wb') as f:
            f.write(cal.to_ical())
        return filepath

    def generate_disbursement_calendar_ics(
        self,
        grant_data: GrantData,
        file_id: str,
        disbursement_interval_days: int = 30,
        disbursement_reminder_days: int = 7,
    ) -> str:
        """Generate a disbursement / payment request calendar.

        Events are drawn from timeline items in the reimbursement/disbursement/
        payment categories.  Each event DESCRIPTION includes a checklist of
        items to complete before submitting the disbursement request.
        """
        filename = f"{file_id}_disbursement_calendar.ics"
        filepath = os.path.join(self.temp_dir, filename)
        cal = Calendar()
        cal.add('prodid', '-//Grant Disbursements//EN//')
        cal.add('version', '2.0')
        cal.add('calscale', 'GREGORIAN')
        cal.add('method', 'PUBLISH')
        cal.add('x-wr-calname', f"Disbursements: {grant_data.grant_title or 'Grant'}")
        cal.add('x-wr-timezone', 'America/Chicago')

        disbursement_categories = {'reimbursement', 'disbursement', 'submission', 'payment'}
        default_checklist = (
            "DISBURSEMENT REQUEST CHECKLIST\n"
            "□ Supporting invoices/receipts gathered and organised\n"
            "□ Expenditures reconciled against approved budget lines\n"
            "□ Budget vs. actuals spreadsheet updated\n"
            "□ All expenses within the approved grant period\n"
            "□ Required approvals (finance/programme director) obtained\n"
            "□ Prior disbursement report filed (if required)\n"
            "□ Submission portal / email confirmed with funder"
        )

        event_counter = 0
        if grant_data.timeline and grant_data.timeline.items:
            for item in grant_data.timeline.items:
                if (item.category or '').lower() not in disbursement_categories:
                    continue
                event_date = self._parse_date_safe(item.date)
                if event_date.year >= 2099:
                    continue

                # Build description with any matching submission requirements
                matching = [
                    req.instructions
                    for req in (grant_data.submission_requirements or [])
                    if req.due_date and req.instructions
                    and abs((self._parse_date_safe(req.due_date) - event_date).days) <= disbursement_reminder_days
                ]
                desc_parts = [item.description]
                if item.amount:
                    desc_parts.append(f"Amount: {item.amount}")
                if matching:
                    desc_parts.append("Funder requirements:\n" + "\n".join(f"- {m}" for m in matching))
                desc_parts.append(default_checklist)

                event = self._build_event(
                    summary=f"[DISBURSEMENT] {item.description}",
                    description='\n\n'.join(desc_parts),
                    start_dt=event_date,
                    end_dt=event_date + timedelta(days=1),
                    all_day=True,
                    alarms=[timedelta(days=-disbursement_reminder_days), timedelta(days=-1)],
                    uid=f"{file_id}-disb-{event_counter}@grantmanagement.local",
                )
                cal.add_component(event)
                event_counter += 1

        with open(filepath, 'wb') as f:
            f.write(cal.to_ical())
        return filepath

    def generate_reporting_calendar_ics(
        self,
        grant_data: GrantData,
        file_id: str,
        disbursement_reminder_days: int = 7,
    ) -> str:
        """Generate a reporting-deadline calendar.

        Draws from both extracted reporting_requirements and timeline items in the
        'report' category.  Each event DESCRIPTION lists the required report elements.
        """
        filename = f"{file_id}_reporting_calendar.ics"
        filepath = os.path.join(self.temp_dir, filename)
        cal = Calendar()
        cal.add('prodid', '-//Grant Reporting Deadlines//EN//')
        cal.add('version', '2.0')
        cal.add('calscale', 'GREGORIAN')
        cal.add('method', 'PUBLISH')
        cal.add('x-wr-calname', f"Reporting: {grant_data.grant_title or 'Grant'}")
        cal.add('x-wr-timezone', 'America/Chicago')

        event_counter = 0

        # Primary source: extracted reporting requirements
        for req in (grant_data.reporting_requirements or []):
            if not req.due_date:
                continue
            event_date = self._parse_date_safe(req.due_date)
            if event_date.year >= 2099:
                continue

            elements = req.required_elements or []
            if elements:
                elements_text = "REQUIRED REPORT ELEMENTS\n" + "\n".join(f"□ {e}" for e in elements)
            else:
                elements_text = (
                    "STANDARD REPORT ELEMENTS\n"
                    "□ Progress against project goals and milestones\n"
                    "□ Budget vs. actuals (with narrative for variances)\n"
                    "□ Beneficiary/participant data (if required)\n"
                    "□ Challenges encountered and adaptations made\n"
                    "□ Planned activities for the next period\n"
                    "□ Supporting documentation (photos, testimonials, etc.)"
                )

            period_label = f" ({req.period})" if req.period else ""
            desc = f"{req.description or 'Progress report'}{period_label}\n\n{elements_text}"

            event = self._build_event(
                summary=f"[REPORT DUE] {req.description or 'Progress Report'}",
                description=desc,
                start_dt=event_date,
                end_dt=event_date + timedelta(days=1),
                all_day=True,
                alarms=[timedelta(days=-14), timedelta(days=-disbursement_reminder_days)],
                uid=f"{file_id}-report-{event_counter}@grantmanagement.local",
            )
            cal.add_component(event)
            event_counter += 1

        # Secondary source: timeline items categorised as 'report' not already covered
        seen_dates = set()
        for item in (grant_data.timeline.items if grant_data.timeline else []):
            if (item.category or '').lower() != 'report':
                continue
            event_date = self._parse_date_safe(item.date)
            if event_date.year >= 2099 or event_date.date() in seen_dates:
                continue
            seen_dates.add(event_date.date())

            desc = (
                f"{item.description}\n\n"
                "STANDARD REPORT ELEMENTS\n"
                "□ Progress against project goals and milestones\n"
                "□ Budget vs. actuals (with narrative for variances)\n"
                "□ Challenges encountered and adaptations made\n"
                "□ Planned activities for the next period"
            )
            event = self._build_event(
                summary=f"[REPORT DUE] {item.description}",
                description=desc,
                start_dt=event_date,
                end_dt=event_date + timedelta(days=1),
                all_day=True,
                alarms=[timedelta(days=-14), timedelta(days=-disbursement_reminder_days)],
                uid=f"{file_id}-report-tl-{event_counter}@grantmanagement.local",
            )
            cal.add_component(event)
            event_counter += 1

        with open(filepath, 'wb') as f:
            f.write(cal.to_ical())
        return filepath

    def generate_summary_docx(self, grant_data: GrantData, file_id: str) -> str:
        """Generate a high-level grant summary Word document.

        Sections:
        1. Grant Overview (key facts table)
        2. Project Purpose & Description
        3. Deliverables & Milestones (timeline items)
        4. Financial Summary (budget breakdown if available)
        5. Reporting Obligations
        6. Data Gaps & Notes
        """
        filename = f"{file_id}_grant_summary.docx"
        filepath = os.path.join(self.temp_dir, filename)
        doc = Document()

        # ── Styles ──────────────────────────────────────────────────────────────
        def _h1(text):
            p = doc.add_heading(text, level=1)
            p.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)  # blue
            return p

        def _h2(text):
            return doc.add_heading(text, level=2)

        def _body(text):
            p = doc.add_paragraph(text)
            p.style = doc.styles['Normal']
            return p

        def _kv_table(rows):
            """Render a two-column key-value table."""
            tbl = doc.add_table(rows=len(rows), cols=2)
            tbl.style = 'Table Grid'
            for i, (key, val) in enumerate(rows):
                tbl.rows[i].cells[0].text = key
                tbl.rows[i].cells[1].text = str(val) if val else "—"
                tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            return tbl

        title = grant_data.grant_title or "Grant"
        org = grant_data.organization_name or "Your Organization"
        funder = grant_data.funder_name or "Funder"

        # Cover heading
        doc.add_heading(f"Grant Summary", 0)
        _body(f"{org}  ·  {funder}  ·  {title}")
        doc.add_paragraph()

        # ── 1. Grant Overview ────────────────────────────────────────────────
        _h1("1. Grant Overview")
        overview_rows = [
            ("Organization", grant_data.organization_name),
            ("Funder", grant_data.funder_name),
            ("Grant Title / Project", grant_data.grant_title),
            ("Grant Amount", f"${grant_data.grant_amount:,.2f}" if grant_data.grant_amount else None),
            ("Grant Period", grant_data.grant_period),
            ("Document Type", (grant_data.document_format or "").replace("_", " ").title()),
        ]
        _kv_table(overview_rows)
        doc.add_paragraph()

        # ── 2. Purpose ───────────────────────────────────────────────────────
        if grant_data.purpose:
            _h1("2. Project Purpose")
            _body(grant_data.purpose)
            doc.add_paragraph()

        # ── 3. Deliverables & Milestones ────────────────────────────────────
        section_num = 3
        timeline_items = (grant_data.timeline.items if grant_data.timeline else []) or []
        _h1(f"{section_num}. Deliverables & Milestones")
        if timeline_items:
            tbl = doc.add_table(rows=1 + len(timeline_items), cols=3)
            tbl.style = 'Table Grid'
            hdr = tbl.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "Date", "Category", "Description"
            for cell in hdr:
                cell.paragraphs[0].runs[0].bold = True
            for i, item in enumerate(timeline_items, 1):
                tbl.rows[i].cells[0].text = item.date or ""
                tbl.rows[i].cells[1].text = (item.category or "").title()
                tbl.rows[i].cells[2].text = item.description or ""
        else:
            _body("No specific milestones were extracted from the grant documents.")
        doc.add_paragraph()

        # ── 4. Financial Summary ─────────────────────────────────────────────
        section_num += 1
        _h1(f"{section_num}. Financial Summary")
        if grant_data.grant_amount:
            _body(f"Total Award: ${grant_data.grant_amount:,.2f}")
        if grant_data.budget and grant_data.budget.items:
            budget_rows = [(item.category, f"${item.amount:,.2f}" if item.amount else "—") for item in grant_data.budget.items]
            _kv_table(budget_rows)
        else:
            _body("Detailed budget breakdown not extracted — refer to the original award documents.")
        doc.add_paragraph()

        # ── 5. Reporting Obligations ─────────────────────────────────────────
        section_num += 1
        _h1(f"{section_num}. Reporting Obligations")
        reporting = grant_data.reporting_requirements or []
        if reporting:
            for req in reporting:
                label = req.description or "Progress Report"
                due = f" — due {req.due_date}" if req.due_date else ""
                period = f" ({req.period})" if req.period else ""
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{label}{period}{due}").bold = False
                if req.required_elements:
                    for el in req.required_elements:
                        sub = doc.add_paragraph(style='List Bullet 2')
                        sub.text = el
        else:
            _body("No specific reporting requirements were found in the grant documents.")
        doc.add_paragraph()

        # ── 6. Data Gaps ─────────────────────────────────────────────────────
        section_num += 1
        gaps = grant_data.data_gaps or []
        if gaps:
            _h1(f"{section_num}. Data Gaps & Recommended Follow-Up")
            for gap in gaps:
                p = doc.add_paragraph(style='List Bullet')
                p.text = gap
            doc.add_paragraph()

        doc.save(filepath)
        return filepath

    def generate_calendar_ics(
        self,
        grant_data: GrantData,
        file_id: str,
        disbursement_interval_days: int = 30,
        disbursement_reminder_days: int = 7,
        meeting_interval_days: int = 14,
    ) -> tuple[str, List[str]]:
        """Legacy combined-calendar generator kept for backward compatibility.
        Delegates to the three separate generators and returns the meeting calendar path.
        """
        self.generate_disbursement_calendar_ics(grant_data, file_id, disbursement_interval_days, disbursement_reminder_days)
        self.generate_reporting_calendar_ics(grant_data, file_id, disbursement_reminder_days)
        meeting_path = self.generate_meeting_calendar_ics(grant_data, file_id, meeting_interval_days)
        discrepancies = self._detect_calendar_discrepancies(grant_data, disbursement_reminder_days)
        return meeting_path, discrepancies

    def _parse_date_safe(self, date_str: str) -> datetime:
        """Safely parse various date formats"""
        from dateutil import parser
        try:
            return parser.parse(date_str)
        except:
            return datetime(2099, 12, 31)
    
    def _extract_account_number(self, category: str) -> str:
        """Extract account number from category name or assign default"""
        # Common account number mappings
        account_map = {
            'personnel': '0005',
            'fringe': '0044',
            'operating': '0100',
            'professional': '0140',
            'travel': '0200',
            'materials': '0300',
            'supplies': '0300',
            'equipment': '0400',
            'indirect': '0801',
            'wages': '0999',
            'stipends': '0050',
        }
        
        category_lower = category.lower()
        for key, account in account_map.items():
            if key in category_lower:
                return account
        
        return '0999'  # Default "Other"
    
    def _extract_amount(self, amount_str: str) -> float:
        """Extract numeric amount from string"""
        if isinstance(amount_str, (int, float)):
            return float(amount_str)
        
        # Remove currency symbols and commas
        amount_str = re.sub(r'[^\d.]', '', str(amount_str))
        try:
            return float(amount_str)
        except:
            return 0.0
    
    def generate_all_documents(self, grant_data: GrantData, file_id: str, options: Dict) -> Dict:
        """Generate all requested documents.

        The options dict may include:
          generate_workplan, generate_budget, generate_report_template,
          generate_calendar, generate_agenda_template  (bool flags)
          disbursement_interval_days, disbursement_reminder_days,
          meeting_interval_days  (int, forwarded to generate_calendar_ics)

        Returns a dict of doc_type -> filepath for downloadable files, plus the
        special key 'calendar_discrepancy' -> List[str] when calendar generation
        is requested.
        """
        generated_files: Dict = {}

        print(f"\n📄 Generating documents for file_id: {file_id}")
        print(f"Options: {options}")

        if options.get('generate_workplan', True):
            try:
                print("  📋 Generating work plan PDF...")
                generated_files['workplan'] = self.generate_workplan_pdf(grant_data, file_id)
                print("  ✓ Work plan created")
            except Exception as e:
                print(f"  ❌ Work plan error: {e}")
                generated_files['workplan_error'] = str(e)

        if options.get('generate_budget', True):
            try:
                print("  💰 Generating budget Excel...")
                generated_files['budget'] = self.generate_budget_excel(grant_data, file_id)
                print("  ✓ Budget created with multiple sheets")
            except Exception as e:
                print(f"  ❌ Budget error: {e}")
                generated_files['budget_error'] = str(e)

        if options.get('generate_report_template', True):
            try:
                print("  📝 Generating report template DOCX...")
                generated_files['report'] = self.generate_report_template_docx(grant_data, file_id)
                print("  ✓ Report template created")
            except Exception as e:
                print(f"  ❌ Report template error: {e}")
                generated_files['report_error'] = str(e)

        if options.get('generate_summary', True):
            try:
                print("  📄 Generating grant summary DOCX...")
                generated_files['summary'] = self.generate_summary_docx(grant_data, file_id)
                print("  ✓ Grant summary created")
            except Exception as e:
                print(f"  ❌ Summary error: {e}")
                generated_files['summary_error'] = str(e)

        disbursement_reminder_days = options.get('disbursement_reminder_days', 7)

        if options.get('generate_meeting_calendar', options.get('generate_calendar', True)):
            try:
                print("  📅 Generating meeting calendar ICS...")
                generated_files['meeting_calendar'] = self.generate_meeting_calendar_ics(
                    grant_data, file_id,
                    meeting_interval_days=options.get('meeting_interval_days', 14),
                )
                print("  ✓ Meeting calendar created")
            except Exception as e:
                print(f"  ❌ Meeting calendar error: {e}")
                generated_files['meeting_calendar_error'] = str(e)

        if options.get('generate_disbursement_calendar', options.get('generate_calendar', True)):
            try:
                print("  📅 Generating disbursement calendar ICS...")
                generated_files['disbursement_calendar'] = self.generate_disbursement_calendar_ics(
                    grant_data, file_id,
                    disbursement_interval_days=options.get('disbursement_interval_days', 30),
                    disbursement_reminder_days=disbursement_reminder_days,
                )
                print("  ✓ Disbursement calendar created")
            except Exception as e:
                print(f"  ❌ Disbursement calendar error: {e}")
                generated_files['disbursement_calendar_error'] = str(e)

        if options.get('generate_reporting_calendar', options.get('generate_calendar', True)):
            try:
                print("  📅 Generating reporting calendar ICS...")
                generated_files['reporting_calendar'] = self.generate_reporting_calendar_ics(
                    grant_data, file_id,
                    disbursement_reminder_days=disbursement_reminder_days,
                )
                discrepancies = self._detect_calendar_discrepancies(grant_data, disbursement_reminder_days)
                generated_files['calendar_discrepancy'] = discrepancies
                print(f"  ✓ Reporting calendar created ({len(discrepancies)} discrepancy notice(s))")
            except Exception as e:
                print(f"  ❌ Reporting calendar error: {e}")
                generated_files['reporting_calendar_error'] = str(e)

        print("✓ Document generation complete\n")

        return generated_files