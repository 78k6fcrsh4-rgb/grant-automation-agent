import os
import uuid
from pypdf import PdfReader
from docx import Document as DocxDocument
from typing import Tuple


def generate_file_id() -> str:
    """Generate unique file ID"""
    return str(uuid.uuid4())


def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from PDF file"""
    try:
        reader = PdfReader(filepath)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"PDF extraction failed: {str(e)}")


def extract_text_from_docx(filepath: str) -> str:
    """Extract text from Word document, including table content."""
    try:
        doc = DocxDocument(filepath)
        parts: list[str] = []

        # Paragraphs preserve reading order within the main body
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

        # Tables are not interleaved with paragraphs in python-docx's paragraph list,
        # so we append them separately.  Each row becomes a pipe-separated line so
        # budget rows like "Premium Assistance | $1,000,000" are extractable.
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    # Deduplicate merged cells (python-docx repeats merged cell text)
                    deduped = [cells[0]]
                    for c in cells[1:]:
                        if c != deduped[-1]:
                            deduped.append(c)
                    parts.append("  |  ".join(deduped))

        return "\n".join(parts).strip()
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")


def extract_text_from_file(filepath: str) -> Tuple[str, str]:
    """
    Extract text from supported file types
    Returns: (text, file_type)
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
    
    ext = os.path.splitext(filepath)[1].lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(filepath), 'pdf'
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(filepath), 'docx'
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def save_uploaded_file(file_content: bytes, filename: str, file_id: str, temp_dir: str = "temp_files") -> str:
    """
    Save uploaded file and return filepath only
    
    Args:
        file_content: The file bytes
        filename: Original filename
        file_id: Pre-generated file ID
        temp_dir: Directory to save files
    
    Returns:
        str: The full filepath where file was saved
    """
    os.makedirs(temp_dir, exist_ok=True)
    ext = os.path.splitext(filename)[1]
    filepath = os.path.join(temp_dir, f"{file_id}{ext}")
    
    with open(filepath, 'wb') as f:
        f.write(file_content)
    
    return filepath